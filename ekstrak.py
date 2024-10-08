import PyPDF2
import re
from docx import Document
import openpyxl
from datetime import datetime

def cari_pdf(file_path, keywords):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)    
        results = []
        
        # Iterasi tiap halaman lalu cari
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text = page.extract_text()
            
            for keyword in keywords:
                matches = re.finditer(keyword, text, re.IGNORECASE)
                
                for match in matches:
                    start = max(0, match.start() - 100)
                    end = min(len(text), match.end() + 100)
                    context = text[start:end]
                    
                    results.append({
                        'keyword': keyword,
                        'page': page_num + 1,
                        'context': context.strip()
                    })
    
    return results

def save_to_word(results, filename):
    doc = Document()
    for result in results:
        doc.add_paragraph(f"Kata kunci '{result['keyword']}' ditemukan di halaman {result['page']}:")
        doc.add_paragraph(result['context'])
        doc.add_paragraph('-' * 50)
    doc.save(filename)

def save_to_excel(results, filename):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Kata Kunci", "Halaman", "Konteks"])
    for result in results:
        ws.append([result['keyword'], result['page'], result['context']])
    wb.save(filename)

file_path = 'perbupPAK.pdf'
keywords = ['Temenggungan', 'Patemon', 'Jatiurip', 'Opo Opo', 'Kamalkuning', 'Tanjungsari', 'Krejengan',
            'Sentong', 'Sumberkatimoho', 'Karangren', 'Rawan', 'Seboro', 'Kedungcaluk', 'Widoro', 'Gebangan',
            'Dawuhan', 'Sokaan', 'Duwuhan']

results = cari_pdf(file_path, keywords)

# nama file sesuai timestamp
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
word_filename = f"hasil_pencarian_{timestamp}.docx"
excel_filename = f"hasil_pencarian_{timestamp}.xlsx"

save_to_word(results, word_filename)
save_to_excel(results, excel_filename)

print(f"Hasil disimpan ke {word_filename} dan {excel_filename}")