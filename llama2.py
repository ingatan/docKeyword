import PyPDF2
import re
from docx import Document
import openpyxl
from datetime import datetime
from tqdm import tqdm
import tabula
import subprocess
import json

def run_llama(prompt):
    command = ['ollama', 'run', 'llama2', prompt]
    result = subprocess.run(command, capture_output=True, text=True)
    return result.stdout.strip()

def analyze_context(context):
    prompt = f"Analyze the following text and provide a brief summary and key points:\n\n{context}\n\nSummary:"
    return run_llama(prompt)

def cari_pdf(file_path, keywords, proximity=5):
    results = []
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        total_pages = len(reader.pages)
        
        for page_num in tqdm(range(total_pages), desc="Memproses halaman"):
            page = reader.pages[page_num]
            text = page.extract_text()
            
            tables = tabula.read_pdf(file_path, pages=page_num+1, multiple_tables=True)
            
            paragraphs = text.split('\n\n')
            for i, paragraph in enumerate(paragraphs):
                for j, keyword in enumerate(keywords):
                    if keyword.lower() in paragraph.lower():
                        context = paragraph.strip()
                        
                        if j < len(keywords) - 1:
                            next_keyword = keywords[j+1]
                            if next_keyword.lower() in paragraph.lower():
                                words = paragraph.lower().split()
                                if abs(words.index(keyword.lower()) - words.index(next_keyword.lower())) <= proximity:
                                    context += f" [Kata kunci berdekatan: {keyword} dan {next_keyword}]"
                        
                        analysis = analyze_context(context)
                        
                        results.append({
                            'keyword': keyword,
                            'page': page_num + 1,
                            'paragraph': i + 1,
                            'context': context,
                            'in_table': any(keyword.lower() in table.to_string().lower() for table in tables),
                            'analysis': analysis,
                        })
    
    return results

def save_to_word(results, filename):
    doc = Document()
    doc.add_heading('Hasil Pencarian Kata Kunci', 0)
    
    for result in results:
        doc.add_paragraph(f"Kata kunci: '{result['keyword']}'")
        doc.add_paragraph(f"Halaman: {result['page']}, Paragraf: {result['paragraph']}")
        doc.add_paragraph(f"Dalam tabel: {'Ya' if result['in_table'] else 'Tidak'}")
        doc.add_paragraph(f"Konteks: {result['context']}")
        doc.add_paragraph(f"Analisis: {result['analysis']}")
        doc.add_paragraph('-' * 50)
    
    doc.add_page_break()
    doc.add_heading('Ringkasan Pencarian', level=1)
    unique_keywords = set(r['keyword'] for r in results)
    for keyword in unique_keywords:
        count = sum(1 for r in results if r['keyword'] == keyword)
        doc.add_paragraph(f"'{keyword}' ditemukan {count} kali")
    
    doc.save(filename)

file_path = 'perbupPAK.pdf'
keywords = ['Temenggungan', 'Patemon', 'Jatiurip', 'Opo Opo', 'Kamalkuning', 'Tanjungsari', 'Krejengan',
            'Sentong', 'Sumberkatimoho', 'Karangren', 'Rawan', 'Seboro', 'Kedungcaluk', 'Widoro', 'Gebangan',
            'Dawuhan', 'Sokaan', 'Duwuhan']

results = cari_pdf(file_path, keywords)

timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
word_filename = f"hasil_pencarian_{timestamp}.docx"

save_to_word(results, word_filename)

print(f"Hasil disimpan ke {word_filename}")